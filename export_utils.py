from docx import Document
import csv
from openpyxl import Workbook


def _normalize_row(r):
    """Convert ANY row into: (text, source, file, score).
       Handles dicts, strings, tuples, lists.
    """

    # ------------------
    # CASE 1: r is a string = pure text
    # ------------------
    if isinstance(r, str):
        return r, "Unknown", "Unknown", "N/A"

    # ------------------
    # CASE 2: r is a list or tuple
    #   Examples:
    #     ("text", {"source": "file"})
    #     ["text only"]
    # ------------------
    if isinstance(r, (list, tuple)):
        text = r[0] if len(r) > 0 else ""
        meta = r[1] if len(r) > 1 and isinstance(r[1], dict) else {}

        source = meta.get("source", "Unknown")
        file_name = meta.get("file", source)
        score = meta.get("score", "N/A")

        return text, source, file_name, score

    # ------------------
    # CASE 3: r is a dict (ideal case)
    # ------------------
    if isinstance(r, dict):
        text = r.get("text", "")

        # meta may or may not exist
        meta = r.get("meta", {})
        if not isinstance(meta, dict):
            meta = {}

        source = r.get("source") or meta.get("source", "Unknown")
        file_name = r.get("file") or source
        score = r.get("score", "N/A")

        return text, source, file_name, score

    # ------------------
    # Unknown structure
    # ------------------
    return str(r), "Unknown", "Unknown", "N/A"


# ======================================================
#  CSV
# ======================================================
def export_to_csv(rows, question, answer, buffer):
    writer = csv.writer(buffer)

    writer.writerow(["Question", question])
    writer.writerow(["Answer", answer])
    writer.writerow([])
    writer.writerow(["Text", "Source", "File", "Score"])

    for r in rows:
        text, source, file_name, score = _normalize_row(r)
        writer.writerow([text, source, file_name, score])


# ======================================================
#  Excel
# ======================================================
def export_to_excel(rows, question, answer, buffer):
    wb = Workbook()
    ws = wb.active
    ws.title = "RAG Export"

    ws.append(["Question", question])
    ws.append(["Answer", answer])
    ws.append([])
    ws.append(["Text", "Source", "File", "Score"])

    for r in rows:
        text, source, file_name, score = _normalize_row(r)
        ws.append([text, source, file_name, score])

    wb.save(buffer)


# ======================================================
#  DOCX
# ======================================================
def export_to_docx(rows, question, answer, buffer):
    doc = Document()

    doc.add_heading("RAG Export", level=1)
    doc.add_paragraph(f"Question: {question}")
    doc.add_paragraph(f"Answer: {answer}")
    doc.add_paragraph("")

    for r in rows:
        text, source, file_name, score = _normalize_row(r)

        doc.add_heading("Retrieved Chunk", level=2)
        doc.add_paragraph(text)

        doc.add_paragraph(f"Source: {source}")
        doc.add_paragraph(f"File: {file_name}")
        doc.add_paragraph(f"Score: {score}")
        doc.add_paragraph("")

    doc.save(buffer)
